from docx import Document

#burada once docx dosyamizi programimiza okutalim
doc = Document("vize_odevi.docx")
#kelimelerimizi tutacagimiz dizi
kelimeler = []
#kutuphane okudugu docx dosyasini kelime kelime olmasi icin dizimize atiyoruz
for p in doc.paragraphs:
    kelimeler.extend(p.text.split())
#ayni kelime olsada buyuk kucuk harf yuzunden program ayni oldugunu anlamayacagundan hepsini buyuk harf yapiyoruz
kelimeler = [k.upper() for k in kelimeler]

#once kelime ve sayilarini tutacagimiz dictionry aciyoruz
kelime_sayilari = {}
#for dongusu ile tek tek kelimeler dizisindeki kelimelere erisiyoruz
for kelime in kelimeler:
    #eger kelimemiz daha once dictiniory e eklenmisse sayisini 1 arttiriyoruz
    if  kelime in kelime_sayilari:
        kelime_sayilari[kelime] += 1
    #eger eklenmemisse sayisini 1 yapiyoruz yani ekliyoruz
    else:
        kelime_sayilari[kelime] = 1

#burada ilk 100 kelimeyi bulmak icin sayisina gore siraladik
#sayiya gore siralanmasi icin key=lambda x: x[1] su yapiyi kullandim
sirala = sorted(kelime_sayilari.items(), key=lambda x: x[1], reverse=True)
#100 kelimeyi yazdiriyoruz
print(sirala[:100])

#kelimelerin basina ve sonuna tirnak isareti ekliyoruz burada buyuk harfe cevirmedim daha once yaptigim icin
tirnakli_kelimeler = [(f'"{kelime}"', sayi) for kelime, sayi in sirala[:100]]
#tirnak isaretli kelimeleride yazdiralim
print(tirnakli_kelimeler)
#once kutuphaneyi kullanarak degiskenimizi dokuman olarak tanimlayalim
ilk_100_kelime = Document()
#dokumanimiza baslik ekledim
ilk_100_kelime.add_heading("Verilen Dosyada En Cok Gecen 100 Kelime ve Sayilari", 1)
#tirnakli sekilde yazilmis elimizdeki 100 kelime dizi icinde oldugundan tek tek erisip kelime ve sayi seklinde dokumana yazdiriyoruz
for kelime, sayi in tirnakli_kelimeler:
    ilk_100_kelime.add_paragraph(f"{kelime} : {sayi}")

#son olarak dokuman bize soylediginiz adda kayit ediliyor
ilk_100_kelime.save("tamamlanmis_vizeödevi.docx")









